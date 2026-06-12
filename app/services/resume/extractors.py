"""简历文件文本提取"""
import io
import logging
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


class PDFExtractor:
    def __init__(self, image_extractor=None):
        self.image_extractor = image_extractor

    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> str:
        if pdfplumber is None:
            logging.info("pdfplumber 未安装，改用 PyMuPDF 文本提取回退")
            return ""
        text_content = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"PDF 解析失败: {e}")
            return ""

    @staticmethod
    def _extract_with_fitz_text(file_path: str) -> str:
        try:
            import fitz
        except ImportError:
            logging.warning("PyMuPDF 未安装，无法执行 PDF 文本回退提取")
            return ""

        text_content = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                page_text = page.get_text("text")
                if page_text:
                    text_content.append(page_text)
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"PyMuPDF 文本提取失败: {e}")
            return ""

    def _extract_with_ocr(self, file_path: str) -> str:
        if not self.image_extractor or not self.image_extractor.available:
            return ""
        try:
            import fitz
            from PIL import Image
        except ImportError:
            logging.warning("PyMuPDF/Pillow 未安装，无法对图片型 PDF 执行 OCR")
            return ""

        text_content = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                page_text = self.image_extractor.extract_image(image)
                if page_text:
                    text_content.append(page_text)
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"PDF OCR 解析失败: {e}")
            return ""

    def extract(self, file_path: str) -> str:
        pdfplumber_text = self._extract_with_pdfplumber(file_path)
        fitz_text = self._extract_with_fitz_text(file_path)
        text = max(
            (pdfplumber_text, fitz_text),
            key=lambda value: len(value.strip()),
            default="",
        )

        # 图片型 PDF 或复杂排版 PDF 常只有很少文本，此时尝试 OCR 回退。
        if len(text.strip()) < 80:
            ocr_text = self._extract_with_ocr(file_path)
            if len(ocr_text.strip()) > len(text.strip()):
                return ocr_text
        return text


class DocxExtractor:
    @staticmethod
    def _extract_with_python_docx(file_path: str) -> str:
        if docx is None:
            logging.error("python-docx 未安装，请运行: pip install python-docx")
            return ""
        text_content = []
        d = docx.Document(file_path)
        for para in d.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        for table in d.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    text_content.append(" | ".join(row_data))
        return "\n".join(text_content)

    @staticmethod
    def _extract_with_mammoth(file_path: str) -> str:
        try:
            import mammoth
        except ImportError:
            return ""

        with open(file_path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        return result.value or ""

    @staticmethod
    def _extract_from_docx_xml(file_path: str) -> str:
        text_content = []
        xml_parts = {
            "word/document.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
            "word/comments.xml",
        }
        try:
            with zipfile.ZipFile(file_path) as archive:
                names = set(archive.namelist())
                xml_parts.update(
                    name for name in names
                    if name.startswith("word/header") or name.startswith("word/footer")
                )
                for part_name in sorted(xml_parts):
                    if part_name not in names:
                        continue
                    root = ET.fromstring(archive.read(part_name))
                    fragments = [
                        node.text.strip()
                        for node in root.iter()
                        if node.tag.endswith("}t") and node.text and node.text.strip()
                    ]
                    if fragments:
                        text_content.append("\n".join(fragments))
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"DOCX XML 回退提取失败: {e}")
            return ""

    @staticmethod
    def extract(file_path: str) -> str:
        if Path(file_path).suffix.lower() == ".doc":
            logging.warning("暂不支持直接解析 .doc，请先转换为 .docx 或 PDF")
            return ""

        candidates = []
        try:
            python_docx_text = DocxExtractor._extract_with_python_docx(file_path)
            if python_docx_text.strip():
                candidates.append(python_docx_text)
        except Exception as e:
            logging.warning(f"python-docx 解析失败，尝试回退: {e}")

        try:
            mammoth_text = DocxExtractor._extract_with_mammoth(file_path)
            if mammoth_text.strip():
                candidates.append(mammoth_text)
        except Exception as e:
            logging.warning(f"mammoth 解析失败: {e}")

        xml_text = DocxExtractor._extract_from_docx_xml(file_path)
        if xml_text.strip():
            candidates.append(xml_text)

        return max(candidates, key=lambda value: len(value.strip()), default="")


class ImageExtractor:
    """图片 OCR 提取（需要 paddleocr，可选依赖）"""
    def __init__(self):
        try:
            from paddleocr import PaddleOCR
            logging.info("正在加载 PaddleOCR 模型...")
            self.ocr = PaddleOCR(use_angle_cls=True, lang="ch")
        except ImportError:
            logging.warning("paddleocr 未安装，图片解析功能不可用")
            self.ocr = None

    @property
    def available(self) -> bool:
        return self.ocr is not None

    def extract(self, file_path: str) -> str:
        if self.ocr is None:
            return ""
        text_content = []
        try:
            result = self.ocr.ocr(file_path, cls=True)
            for idx in range(len(result)):
                res = result[idx]
                if res is not None:
                    for line in res:
                        text_content.append(line[1][0])
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"图片 OCR 解析失败: {e}")
            return ""

    def extract_image(self, image) -> str:
        if self.ocr is None:
            return ""
        text_content = []
        try:
            try:
                import numpy as np
                ocr_input = np.array(image)
            except ImportError:
                ocr_input = image
            result = self.ocr.ocr(ocr_input, cls=True)
            for res in result:
                if res is not None:
                    for line in res:
                        text_content.append(line[1][0])
            return "\n".join(text_content)
        except Exception as e:
            logging.error(f"内存图片 OCR 解析失败: {e}")
            return ""
